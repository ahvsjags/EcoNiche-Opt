from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
PAPER = ROOT / "paper"
SRC = PAPER / "econiche_opt_cover_letter_en_v1_20260509.md"
OUT = PAPER / "econiche_opt_cover_letter_en_v1_20260509.docx"


def set_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    title = styles["Title"]
    title.font.name = "Arial"
    title.font.size = Pt(18)
    title.font.bold = True
    title.font.color.rgb = RGBColor(20, 55, 90)


def add_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Inches(0)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.add_run(text)


def main() -> int:
    text = SRC.read_text(encoding="utf-8").splitlines()
    doc = Document()
    set_styles(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    title = text[0].lstrip("# ").strip()
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(title)

    buffer: list[str] = []
    for line in text[1:]:
        stripped = line.strip()
        if not stripped:
            if buffer:
                add_paragraph(doc, " ".join(buffer))
                buffer = []
            continue
        buffer.append(stripped)
    if buffer:
        add_paragraph(doc, " ".join(buffer))

    doc.save(OUT)
    print(OUT)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

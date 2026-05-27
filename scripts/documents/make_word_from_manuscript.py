from __future__ import annotations

import csv
import re
import zipfile
from pathlib import Path

import matplotlib

matplotlib.use("Agg", force=True)
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
PAPER = ROOT / "paper"
FIG_DIR = ROOT / "figures" / "article"
TABLE_DIR = ROOT / "tables" / "article"
OUT_DIR = PAPER

MAIN_MD = PAPER / "econiche_opt_manuscript_cn_v1_20260508.md"
SUPP_MD = PAPER / "econiche_opt_supporting_cn_v1_20260508.md"
MAIN_EN_MD = PAPER / "econiche_opt_manuscript_en_v1_20260509.md"
SUPP_EN_MD = PAPER / "econiche_opt_supporting_en_v1_20260509.md"
MAIN_DOCX = OUT_DIR / "econiche_opt_manuscript_cn_v1_20260508.docx"
SUPP_DOCX = OUT_DIR / "econiche_opt_supporting_cn_v1_20260508.docx"
MAIN_EN_DOCX = OUT_DIR / "econiche_opt_manuscript_en_v1_20260509.docx"
SUPP_EN_DOCX = OUT_DIR / "econiche_opt_supporting_en_v1_20260509.docx"
EQUATION_DIR = PAPER / "communications_medicine_submission" / "equations_600dpi"

AUTHOR_TEXT = (
    "Pengyuan Xu1,3, Guang Yang2,3, Moyan Li3*\n"
    "1 Department of Materials Science and Engineering, Monash University, Clayton, VIC 3800, Australia.\n"
    "2 School of Economics and Management, China University of Mining and Technology, Xuzhou 221116, Jiangsu, China.\n"
    "3 Hong Kong University of Science and Technology (Guangzhou), Guangzhou 510000, China.\n"
    "*Correspondence: moyanli@hkust-gz.edu.cn"
)

MAIN_FIGURES_HIRES = {
    1: "fig1_study_design_benchmark.png",
    2: "fig2_model_optimizer.png",
    3: "fig3_primary_melanoma_performance.png",
    4: "fig4_robustness_ablation_claims.png",
    5: "fig5_external_panel_rescue.png",
    6: "fig6_mechanism_perturbation.png",
    7: "fig7_translation_package.png",
}

EQUATIONS = {
    1: [r"z_{ig}=\Phi^{-1}\left(\operatorname{rank}_{ig}(X_{i\cdot})\right)"],
    2: [r"d_g=\operatorname{sign}\left\{\operatorname{cor}_{T}(z_g,y)\right\},\quad d_g\in\{-1,+1\}"],
    3: [r"M_{iq}=\frac{1}{\sqrt{|A_{iq}|}}\sum_{g\in A_{iq}} d_g z_{ig}"],
    4: [
        r"I_{i,qr}=Z\!\left\{\frac{1}{|E_{qr}|}\sum_{(g,h,c)\in E_{qr}}(d_gz_{ig})(d_hz_{ih})A^*_{iq}A^*_{ir}\right\}",
    ],
    5: [r"\widehat{p}_i=\Pr(y_i=1\mid F_i)=\sigma\left(\alpha+\beta^\top\widetilde{F}_i\right)"],
    6: [
        r"J(\theta)=\overline{\mathrm{AUROC}}-\rho\,\mathrm{sd}(\mathrm{AUROC})",
        r"+0.10\,\overline{\mathrm{AUPRC}}+0.05\,\overline{\mathrm{BA}}",
        r"-0.15\,\overline{\mathrm{ECE}}+B(\theta)-P(\theta)",
    ],
    7: [
        r"B(\theta)=0.08C_{\mathrm{cell}}+0.04C_{\mathrm{pathway}}",
        r"+0.04C_{\mathrm{network}}+0.04C_{\mathrm{LR}}",
        r"+0.08C_{\mathrm{direction}}",
    ],
    8: [
        r"P(\theta)=0.04P_{\mathrm{size}}+0.08P_{\mathrm{batch}}",
        r"+0.05P_{\mathrm{redundancy}}+0.05P_{\mathrm{therapy}}",
    ],
    9: [r"\theta^*=\arg\max_{\theta\in\Omega_T}J(\theta)"],
    10: [
        r"\tau^*=\arg\max_{\tau\in\{\widehat{p}_i:i\in T_h\}}",
        r"\mathrm{BA}\left(y_i,\mathrm{I}(\widehat{p}_i\geq\tau)\right)",
    ],
    11: [r"\mathrm{NB}(\tau)=\frac{\mathrm{TP}(\tau)}{n}-\frac{\mathrm{FP}(\tau)}{n}\frac{\tau}{1-\tau}"],
    12: [r"\Delta_s=\mathrm{AUROC}_{\mathrm{EcoNiche},s}-\frac{1}{8}\sum_{b=1}^{8}\mathrm{AUROC}_{b,s}"],
}


SUPP_FIGURES = {
    1: "suppfig1_cohort_curation.png",
    2: "suppfig2_platform_gene_coverage.png",
    3: "suppfig3_benchmark_detail.png",
    4: "suppfig4_full_baseline_comparison.png",
    5: "suppfig5_endpoint_sensitivity.png",
    6: "suppfig6_ablation_optimizer.png",
    7: "suppfig7_external_validation_expanded.png",
    8: "suppfig8_pd1_rescue.png",
    9: "suppfig9_single_cell_mechanism.png",
    10: "suppfig10_reproducibility_package.png",
}

SUPP_TABLES = {
    int(path.name.split("_")[2]): path
    for path in TABLE_DIR.glob("supp_table_*.tsv")
    if path.name.split("_")[2].isdigit()
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=80, bottom=80, end=80) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_east_asia_font(run, font_name: str) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def latex_expr_to_plain(expr: str) -> str:
    """Convert lightweight manuscript LaTeX math into readable Word text."""
    replacements = {
        r"\Phi": "Phi",
        r"\theta": "theta",
        r"\rho": "rho",
        r"\beta": "beta",
        r"\alpha": "alpha",
        r"\tau": "tau",
        r"\sigma": "sigmoid",
        r"\Pr": "Pr",
        r"\sum": "sum",
        r"\cap": " intersect ",
        r"\cup": " union ",
        r"\in": " in ",
        r"\setminus": " minus ",
        r"\arg\max": "argmax",
        r"\top": "T",
        r"\hat": "hat",
        r"\tilde": "tilde",
        r"\cdot": ".",
        r"\quad": "; ",
        r"\,": "",
    }
    out = expr.replace(r"\left", "").replace(r"\right", "")
    out = re.sub(r"\\mathrm\{([^{}]+)\}", lambda m: m.group(1).replace(r"\ ", " "), out)
    out = re.sub(r"\\mathcal\{([^{}]+)\}", r"\1", out)
    out = re.sub(r"\\mathbf\{([^{}]+)\}", r"\1", out)
    out = re.sub(r"_\{([^{}]+)\}", r"_\1", out)
    out = re.sub(r"\^\{([^{}]+)\}", r"^\1", out)
    out = re.sub(r"\\overline\{([^{}]+)\}", r"mean(\1)", out)
    out = re.sub(r"\\frac\{([^{}]+)\}\{([^{}]+)\}", r"(\1)/(\2)", out)
    for old, new in replacements.items():
        out = out.replace(old, new)
    out = out.replace(r"\{", "{").replace(r"\}", "}")
    out = re.sub(r"\\([A-Za-z]+)", r"\1", out)
    out = out.replace("{", "").replace("}", "")
    out = re.sub(r"\s+", " ", out).strip()
    return out


def latex_to_word_text(text: str) -> str:
    text = text.replace(r"\[", "").replace(r"\]", "")
    text = re.sub(r"\\\((.*?)\\\)", lambda m: latex_expr_to_plain(m.group(1)), text)
    if "\\" in text:
        text = latex_expr_to_plain(text)
    return text


def set_paragraph_border(paragraph, side: str, color: str, size: str = "8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), size)
    border.set(qn("w:space"), "1")
    border.set(qn("w:color"), color)
    p_bdr.append(border)


def set_section_columns(section, columns: int = 1, space: int = 360) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(columns))
    cols.set(qn("w:space"), str(space))


def configure_styles(doc: Document, mode: str = "main") -> None:
    styles = doc.styles
    is_supp = mode == "supp"
    body_font = "Arial" if is_supp else "Times New Roman"
    body_size = 9.0 if is_supp else 8.5
    normal = styles["Normal"]
    normal.font.name = body_font
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(body_size)

    for style_name, size, color in [
        ("Heading 1", 14.0 if is_supp else 12.0, "000000"),
        ("Heading 2", 12.0 if is_supp else 9.4, "000000"),
        ("Heading 3", 10.0 if is_supp else 8.8, "000000"),
    ]:
        style = styles[style_name]
        style.font.name = body_font
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)

    title = styles["Title"]
    title.font.name = body_font
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title.font.size = Pt(26 if not is_supp else 16)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string("000000")


def add_runs_from_markdown(
    paragraph,
    text: str,
    bold_default: bool = False,
    default_font: str = "Times New Roman",
) -> None:
    # Minimal inline markdown handling for **bold** and `code`.
    text = latex_to_word_text(text)
    tokens = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for token in tokens:
        if not token:
            continue
        bold = bold_default
        font = default_font
        if token.startswith("**") and token.endswith("**"):
            token = token[2:-2]
            bold = True
        elif token.startswith("`") and token.endswith("`"):
            token = token[1:-1]
            font = "Consolas"
        run = paragraph.add_run(token)
        run.bold = bold
        set_east_asia_font(run, font if font != "Consolas" else "Microsoft YaHei")
        if font == "Consolas":
            run.font.name = "Consolas"


def add_markdown_table(doc: Document, rows: list[str]) -> None:
    parsed = []
    for line in rows:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        parsed.append(cells)
    # Drop markdown separator row.
    parsed = [r for r in parsed if not all(re.match(r"^:?-{3,}:?$", c) for c in r)]
    if not parsed:
        return
    table = doc.add_table(rows=len(parsed), cols=max(len(r) for r in parsed))
    table.style = "Table Grid"
    table.autofit = True
    for i, row in enumerate(parsed):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            add_runs_from_markdown(p, value, bold_default=(i == 0))
            set_cell_margins(cell)
            if i == 0:
                set_cell_shading(cell, "D9EAF7")
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
    doc.add_paragraph()


def add_section_break(
    doc: Document,
    landscape: bool = False,
    start_type=WD_SECTION.NEW_PAGE,
    columns: int = 1,
) -> None:
    section = doc.add_section(start_type)
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        section.top_margin = Inches(0.45)
        section.bottom_margin = Inches(0.45)
        section.left_margin = Inches(0.45)
        section.right_margin = Inches(0.45)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
    set_section_columns(section, columns)


def add_tsv_table(doc: Document, table_num: int) -> None:
    table_path = SUPP_TABLES.get(table_num)
    if table_path is None or not table_path.exists():
        note = doc.add_paragraph()
        add_runs_from_markdown(note, f"[Supplementary Table {table_num} file not found]")
        return

    with table_path.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.reader(handle, delimiter="\t"))
    if not rows:
        note = doc.add_paragraph()
        add_runs_from_markdown(note, f"[Supplementary Table {table_num} is empty: {table_path.name}]")
        return

    n_cols = max(len(row) for row in rows)
    n_rows = len(rows)
    font_size = 5.0 if n_cols >= 18 else 5.7 if n_cols >= 10 else 6.8

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(3)
    add_runs_from_markdown(note, f"Full tabulation ({n_rows - 1} rows x {n_cols} columns).", bold_default=False)

    table = doc.add_table(rows=1, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    def fill_row(word_row, values: list[str], is_header: bool = False) -> None:
        for col_idx in range(n_cols):
            value = values[col_idx] if col_idx < len(values) else ""
            cell = word_row.cells[col_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cell, top=35, start=35, bottom=35, end=35)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            run.font.name = "Arial"
            run.font.size = Pt(font_size)
            run.bold = is_header
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            if is_header:
                set_cell_shading(cell, "D9EAF7")

    fill_row(table.rows[0], rows[0], is_header=True)
    set_repeat_table_header(table.rows[0])
    for values in rows[1:]:
        fill_row(table.add_row(), values)
    doc.add_paragraph()


def add_image_if_available(doc: Document, image_path: Path, width_inches: float = 6.5) -> None:
    if not image_path.exists():
        note = doc.add_paragraph()
        add_runs_from_markdown(note, f"[Image file not found: {image_path.name}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))


def render_equation_image(eq_num: int) -> Path:
    EQUATION_DIR.mkdir(parents=True, exist_ok=True)
    out = EQUATION_DIR / f"equation_{eq_num:02d}.png"
    lines = EQUATIONS[eq_num]
    fig_height = 0.38 + 0.22 * (len(lines) - 1)
    fig = plt.figure(figsize=(3.25, fig_height), dpi=600)
    ax = fig.add_axes([0, 0, 1, 1])
    ax.axis("off")
    if len(lines) == 1:
        y_positions = [0.50]
    elif len(lines) == 2:
        y_positions = [0.64, 0.34]
    else:
        y_positions = [0.72, 0.50, 0.28][: len(lines)]
    font_size = 8.8 if eq_num in {4, 6, 7, 8, 10, 12} else 9.4
    for line, y in zip(lines, y_positions):
        ax.text(0.47, y, f"${line}$", ha="center", va="center", fontsize=font_size)
    ax.text(0.985, 0.50, f"({eq_num})", ha="right", va="center", fontsize=8.6)
    fig.savefig(out, dpi=600, transparent=True, bbox_inches="tight", pad_inches=0.02)
    plt.close(fig)
    return out


def add_equation_if_formula_line(doc: Document, line: str) -> bool:
    match = re.match(r"^Formula\s+(\d+):\s*`.*`\s*$", line.strip())
    if not match:
        return False
    eq_num = int(match.group(1))
    if eq_num not in EQUATIONS:
        return False
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_together = True
    run = p.add_run()
    run.add_picture(str(render_equation_image(eq_num)), width=Inches(3.05))
    return True


def add_main_figure(doc: Document, line: str) -> bool:
    match = re.match(r"^\*\*((?:图|Figure)\s*(\d+)\s*\|[^*]+?)\*\*\s*(.*)$", line)
    if not match:
        return False

    title_text = match.group(1)
    fig_num = int(match.group(2))
    legend_text = match.group(3).strip()
    fig_name = MAIN_FIGURES_HIRES.get(fig_num, f"fig{fig_num}.png")

    title = doc.add_paragraph(style="Heading 3")
    add_runs_from_markdown(title, f"**{title_text}**", bold_default=True)

    image_width = 6.5
    add_image_if_available(doc, FIG_DIR / fig_name, width_inches=image_width)

    if legend_text:
        caption = doc.add_paragraph()
        caption.paragraph_format.first_line_indent = Inches(0)
        caption.paragraph_format.line_spacing = 1.0
        caption.paragraph_format.space_after = Pt(8)
        add_runs_from_markdown(caption, legend_text)
        caption_size = 8.1 if fig_num == 7 else 8.8
        for run in caption.runs:
            run.font.size = Pt(caption_size)
            run.font.color.rgb = RGBColor.from_string("222222")
    doc.add_page_break()
    return True


def add_main_figure_reference(doc: Document, line: str) -> bool:
    match = re.match(r"^\*\*((?:(?:Figure)|(?:图))\s*(\d+)\s*\|[^*]+?)\*\*\s*(.*)$", line)
    if not match:
        return False
    title_text = match.group(1)
    fig_num = int(match.group(2))
    legend_text = match.group(3).strip()
    fig_name = MAIN_FIGURES_HIRES.get(fig_num, f"fig{fig_num}.png")

    title = doc.add_paragraph(style="Heading 3")
    title.paragraph_format.space_before = Pt(2)
    title.paragraph_format.space_after = Pt(4)
    add_runs_from_markdown(title, f"**{title_text}**", bold_default=True, default_font="Arial")

    add_image_if_available(doc, FIG_DIR / fig_name, width_inches=7.0)

    if legend_text:
        caption = doc.add_paragraph()
        caption.paragraph_format.first_line_indent = Inches(0)
        caption.paragraph_format.line_spacing = 1.0
        caption.paragraph_format.space_after = Pt(7)
        add_runs_from_markdown(caption, legend_text, default_font="Arial")
        for run in caption.runs:
            run.font.size = Pt(7.8)
            run.font.color.rgb = RGBColor.from_string("222222")
    doc.add_page_break()
    return True


def maybe_add_figure_after_caption(doc: Document, line: str, mode: str) -> None:
    if mode == "main":
        match = re.match(r"^\*\*图\s*(\d+)\s*\|", line)
        if match:
            fig_num = int(match.group(1))
            fig_name = MAIN_FIGURES_HIRES.get(fig_num, f"fig{fig_num}.png")
            add_image_if_available(doc, FIG_DIR / fig_name)
            doc.add_page_break()
    elif mode == "supp":
        match = re.match(r"^\*\*Supplementary Figure\s+(\d+)\s*\|", line)
        if match:
            fig_num = int(match.group(1))
            add_image_if_available(doc, FIG_DIR / SUPP_FIGURES[fig_num])


def add_cover(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.style = "Title"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runs_from_markdown(p, title, bold_default=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(subtitle)
    run.italic = True
    set_east_asia_font(run, "Microsoft YaHei")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string("666666")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("EcoNiche-Opt manuscript draft")
    set_east_asia_font(run, "Microsoft YaHei")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string("777777")
    doc.add_page_break()


def extract_markdown_section(lines: list[str], heading: str) -> tuple[str, set[int]]:
    start = None
    for idx, line in enumerate(lines):
        if line.strip() == heading:
            start = idx
            break
    if start is None:
        return "", set()
    end = len(lines)
    for idx in range(start + 1, len(lines)):
        if lines[idx].startswith("## ") and lines[idx].strip() != heading:
            end = idx
            break
    body = "\n".join(line.rstrip() for line in lines[start + 1 : end]).strip()
    return body, set(range(start, end))


def add_main_frontmatter(doc: Document, title: str, abstract_text: str) -> None:
    red_rule = doc.add_paragraph()
    red_rule.paragraph_format.space_before = Pt(0)
    red_rule.paragraph_format.space_after = Pt(2)
    red_run = red_rule.add_run("_" * 150)
    red_run.font.name = "Arial"
    red_run.font.size = Pt(4)
    red_run.font.color.rgb = RGBColor.from_string("E2231A")

    brand = doc.add_paragraph()
    brand.paragraph_format.space_before = Pt(0)
    brand.paragraph_format.space_after = Pt(24)
    run = brand.add_run("communications medicine")
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(18)
    run.bold = True

    meta = doc.add_table(rows=1, cols=2)
    meta.autofit = True
    meta.cell(0, 0).text = ""
    meta.cell(0, 1).text = ""
    left = meta.cell(0, 0).paragraphs[0]
    left.paragraph_format.space_after = Pt(0)
    add_runs_from_markdown(left, "Article", bold_default=True)
    right = meta.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.paragraph_format.space_after = Pt(0)
    add_runs_from_markdown(right, "Article | EcoNiche-Opt locked immune-ecology score")
    for cell in meta.row_cells(0):
        set_cell_margins(cell, top=0, start=0, bottom=0, end=0)

    title_p = doc.add_paragraph(style="Title")
    title_p.paragraph_format.space_before = Pt(3)
    title_p.paragraph_format.space_after = Pt(18)
    set_paragraph_border(title_p, "top", "777777", size="4")
    set_paragraph_border(title_p, "bottom", "777777", size="4")
    add_runs_from_markdown(title_p, title, bold_default=True)

    authors = doc.add_paragraph()
    authors.paragraph_format.space_after = Pt(10)
    add_runs_from_markdown(authors, AUTHOR_TEXT, bold_default=True)
    for run in authors.runs:
        run.font.size = Pt(9.2)

    abstract_head = doc.add_paragraph()
    abstract_head.paragraph_format.space_after = Pt(1)
    add_runs_from_markdown(abstract_head, "Abstract", bold_default=True)
    for run in abstract_head.runs:
        run.font.size = Pt(10.0)

    abstract_blocks = [block.strip() for block in re.split(r"\n\s*\n", abstract_text) if block.strip()]
    for idx, block in enumerate(abstract_blocks):
        abstract = doc.add_paragraph()
        abstract.paragraph_format.first_line_indent = Inches(0)
        abstract.paragraph_format.line_spacing = 1.0
        abstract.paragraph_format.space_after = Pt(3 if idx < len(abstract_blocks) - 1 else 8)
        add_runs_from_markdown(abstract, block)
        for run in abstract.runs:
            run.font.size = Pt(8.6)


def add_supp_frontmatter(doc: Document, subtitle: str) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(32)
    title.paragraph_format.space_after = Pt(20)
    run = title.add_run("Supplementary Information")
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(16)

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(12)
    add_runs_from_markdown(sub, subtitle, bold_default=True, default_font="Arial")
    for run in sub.runs:
        run.font.size = Pt(14)

    authors = doc.add_paragraph()
    authors.paragraph_format.space_after = Pt(28)
    add_runs_from_markdown(authors, AUTHOR_TEXT, default_font="Arial")
    for run in authors.runs:
        run.font.size = Pt(11.5)


def save_docx_with_fallback(doc: Document, out_path: Path) -> Path:
    try:
        doc.save(out_path)
        add_no_image_compression_setting(out_path)
        return out_path
    except PermissionError:
        fallback = out_path.with_name(f"{out_path.stem}_polished{out_path.suffix}")
        doc.save(fallback)
        add_no_image_compression_setting(fallback)
        return fallback


def add_no_image_compression_setting(docx_path: Path) -> None:
    with zipfile.ZipFile(docx_path, "r") as zin:
        parts = {name: zin.read(name) for name in zin.namelist()}
    settings_name = "word/settings.xml"
    settings = parts.get(settings_name)
    if not settings:
        return
    text = settings.decode("utf-8")
    if "doNotCompressPictures" not in text:
        text = text.replace("</w:settings>", '<w:doNotCompressPictures/></w:settings>')
    if "defaultImageDpi" not in text:
        text = text.replace("</w:settings>", '<w:defaultImageDpi w:val="600"/></w:settings>')
    parts[settings_name] = text.encode("utf-8")
    tmp = docx_path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(tmp, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for name, data in parts.items():
            zout.writestr(name, data)
    tmp.replace(docx_path)


def convert_markdown_to_docx(md_path: Path, out_path: Path, mode: str) -> Path:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    title = next((ln[2:].strip() for ln in lines if ln.startswith("# ")), md_path.stem)

    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    cover_subtitle = "Author information pending" if "_en_" in md_path.name else "作者信息待补"
    add_cover(doc, title, cover_subtitle)

    table_buffer: list[str] = []
    pending_supp_table: int | None = None
    supp_tables_landscape_started = False
    for raw in lines:
        line = raw.rstrip()
        if line.strip() in {r"\[", r"\]", "[作者信息待补]", "[Author information pending]"}:
            continue
        if line.startswith("|") and line.endswith("|"):
            table_buffer.append(line)
            continue
        if table_buffer:
            add_markdown_table(doc, table_buffer)
            table_buffer = []
        if not line.strip():
            continue

        if pending_supp_table is not None and (
            line.startswith("## ") or line.startswith("### ") or re.match(r"^\*\*Supplementary Table\s+\d+\s*\|", line)
        ):
            add_tsv_table(doc, pending_supp_table)
            pending_supp_table = None

        if line.startswith("# "):
            continue
        if mode == "main" and add_equation_if_formula_line(doc, line):
            continue
        if line.startswith("## "):
            if mode == "supp" and line[3:].strip() == "Supplementary Tables" and not supp_tables_landscape_started:
                add_section_break(doc, landscape=True)
                supp_tables_landscape_started = True
            p = doc.add_paragraph(style="Heading 1")
            add_runs_from_markdown(p, line[3:].strip(), bold_default=True)
        elif line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_runs_from_markdown(p, line[4:].strip(), bold_default=True)
        elif mode == "main" and (line.startswith("**图 ") or line.startswith("**Figure ")) and add_main_figure(doc, line.strip()):
            continue
        elif line.startswith("**图 ") or line.startswith("**Figure ") or line.startswith("**Supplementary Figure ") or (line.startswith("**") and line.endswith("**")):
            if mode == "supp" and re.match(r"^\*\*Supplementary Figure\s+10\s*\|", line.strip()):
                doc.add_page_break()
            p = doc.add_paragraph(style="Heading 3")
            add_runs_from_markdown(p, line.strip(), bold_default=True)
            maybe_add_figure_after_caption(doc, line.strip(), mode)
            table_match = re.match(r"^\*\*Supplementary Table\s+(\d+)\s*\|", line.strip())
            if mode == "supp" and table_match:
                pending_supp_table = int(table_match.group(1))
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs_from_markdown(p, line[2:].strip())
            if pending_supp_table is not None:
                add_tsv_table(doc, pending_supp_table)
                pending_supp_table = None
        else:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.22)
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(5)
            add_runs_from_markdown(p, line.strip())
            if pending_supp_table is not None:
                add_tsv_table(doc, pending_supp_table)
                pending_supp_table = None
    if table_buffer:
        add_markdown_table(doc, table_buffer)
    if pending_supp_table is not None:
        add_tsv_table(doc, pending_supp_table)

    doc.core_properties.title = title
    doc.core_properties.subject = "EcoNiche-Opt manuscript draft"
    doc.core_properties.author = "EcoNiche-Opt"
    return save_docx_with_fallback(doc, out_path)


def convert_markdown_to_docx_reference(md_path: Path, out_path: Path, mode: str) -> Path:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    title = next((ln[2:].strip() for ln in lines if ln.startswith("# ")), md_path.stem)
    default_font = "Arial" if mode == "supp" else "Times New Roman"
    skip_indices: set[int] = set()

    doc = Document()
    configure_styles(doc, mode)
    section = doc.sections[0]
    if mode == "main":
        section.page_width = Inches(8.27)
        section.page_height = Inches(10.98)
        section.top_margin = Inches(0.45)
        section.bottom_margin = Inches(0.45)
        section.left_margin = Inches(0.52)
        section.right_margin = Inches(0.52)
        set_section_columns(section, 1)
        abstract_text, abstract_skip = extract_markdown_section(lines, "## Abstract")
        skip_indices |= abstract_skip
        add_main_frontmatter(doc, title, abstract_text)
    else:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(0.72)
        section.bottom_margin = Inches(0.72)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        set_section_columns(section, 1)
        supp_title = next((ln[3:].strip() for ln in lines if ln.startswith("## ")), title)
        add_supp_frontmatter(doc, supp_title)
        for idx, line in enumerate(lines):
            if line.startswith("# ") or line.strip() == f"## {supp_title}":
                skip_indices.add(idx)

    table_buffer: list[str] = []
    pending_supp_table: int | None = None
    supp_tables_landscape_started = False
    main_two_column_started = False
    main_one_column_figures_started = False

    for idx, raw in enumerate(lines):
        if idx in skip_indices:
            continue
        line = raw.rstrip()
        if line.strip() in {r"\[", r"\]", "[Author information pending]"}:
            continue
        if line.strip() in set(AUTHOR_TEXT.splitlines()):
            continue
        if line.startswith("|") and line.endswith("|"):
            table_buffer.append(line)
            continue
        if table_buffer:
            add_markdown_table(doc, table_buffer)
            table_buffer = []
        if not line.strip():
            continue

        if pending_supp_table is not None and (
            line.startswith("## ") or line.startswith("### ") or re.match(r"^\*\*Supplementary Table\s+\d+\s*\|", line)
        ):
            add_tsv_table(doc, pending_supp_table)
            pending_supp_table = None

        if line.startswith("# "):
            continue
        if mode == "main" and add_equation_if_formula_line(doc, line):
            continue
        if line.startswith("## "):
            heading_text = line[3:].strip()
            if mode == "main" and not main_two_column_started and heading_text in {"Introduction", "Results"}:
                add_section_break(doc, start_type=WD_SECTION.CONTINUOUS, columns=2)
                main_two_column_started = True
            if mode == "main" and heading_text == "Figure Legends" and not main_one_column_figures_started:
                add_section_break(doc, start_type=WD_SECTION.NEW_PAGE, columns=1)
                main_one_column_figures_started = True
            if mode == "supp" and heading_text == "Supplementary Tables" and not supp_tables_landscape_started:
                add_section_break(doc, landscape=True)
                supp_tables_landscape_started = True
            p = doc.add_paragraph(style="Heading 1")
            p.paragraph_format.space_before = Pt(7 if mode == "main" else 12)
            p.paragraph_format.space_after = Pt(3 if mode == "main" else 6)
            add_runs_from_markdown(p, heading_text, bold_default=True, default_font=default_font)
        elif line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            p.paragraph_format.space_before = Pt(5 if mode == "main" else 8)
            p.paragraph_format.space_after = Pt(2)
            add_runs_from_markdown(p, line[4:].strip(), bold_default=True, default_font=default_font)
        elif mode == "main" and (line.startswith("**Figure ") or line.startswith("**图 ")) and add_main_figure_reference(doc, line.strip()):
            continue
        elif line.startswith("**Figure ") or line.startswith("**Supplementary Figure ") or (line.startswith("**") and line.endswith("**")):
            if mode == "supp" and re.match(r"^\*\*Supplementary Figure\s+10\s*\|", line.strip()):
                doc.add_page_break()
            p = doc.add_paragraph(style="Heading 3")
            add_runs_from_markdown(p, line.strip(), bold_default=True, default_font=default_font)
            maybe_add_figure_after_caption(doc, line.strip(), mode)
            table_match = re.match(r"^\*\*Supplementary Table\s+(\d+)\s*\|", line.strip())
            if mode == "supp" and table_match:
                pending_supp_table = int(table_match.group(1))
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3 if mode == "main" else 5)
            add_runs_from_markdown(p, line[2:].strip(), default_font=default_font)
            if pending_supp_table is not None:
                add_tsv_table(doc, pending_supp_table)
                pending_supp_table = None
        else:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.16 if mode == "main" else 0.22)
            p.paragraph_format.line_spacing = 1.0 if mode == "main" else 1.12
            p.paragraph_format.space_after = Pt(2 if mode == "main" else 5)
            add_runs_from_markdown(p, line.strip(), default_font=default_font)
            if pending_supp_table is not None:
                add_tsv_table(doc, pending_supp_table)
                pending_supp_table = None

    if table_buffer:
        add_markdown_table(doc, table_buffer)
    if pending_supp_table is not None:
        add_tsv_table(doc, pending_supp_table)

    doc.core_properties.title = title
    doc.core_properties.subject = "EcoNiche-Opt manuscript"
    doc.core_properties.author = "EcoNiche-Opt"
    return save_docx_with_fallback(doc, out_path)


def main() -> None:
    outputs = []
    if MAIN_EN_MD.exists():
        outputs.append(convert_markdown_to_docx_reference(MAIN_EN_MD, MAIN_EN_DOCX, "main"))
    if SUPP_EN_MD.exists():
        outputs.append(convert_markdown_to_docx_reference(SUPP_EN_MD, SUPP_EN_DOCX, "supp"))
    for output in outputs:
        print(output)


if __name__ == "__main__":
    main()

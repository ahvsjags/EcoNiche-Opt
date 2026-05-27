from __future__ import annotations

import re
from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from openpyxl import load_workbook
from openpyxl.styles import Font, PatternFill
from openpyxl.utils import get_column_letter


ROOT = Path(__file__).resolve().parents[2]
PAPER_DIR = ROOT / "paper"
TABLE_DIR = ROOT / "tables" / "article"
FIGURE_DIR = ROOT / "figures" / "article"
SUBMISSION_DIR = PAPER_DIR / "communications_medicine_submission"
SOURCE_DATA = SUBMISSION_DIR / "Source_Data.xlsx"
CHECKLISTS = SUBMISSION_DIR / "STROBE_TRIPOD_REMARK_checklists.xlsx"
REPORTING_SUMMARY = SUBMISSION_DIR / "NaturePortfolio_Reporting_Summary_EcoNicheOpt.docx"
READINESS = SUBMISSION_DIR / "communications_medicine_submission_readiness.md"

RELEASE_TAG = "v0.3.1-jtm-20260527"
REPO = "https://github.com/ahvsjags/EcoNiche-Opt"
ARCHIVE = f"{REPO}/archive/refs/tags/{RELEASE_TAG}.zip"


FIGURE_SOURCE_MAP = [
    ("Figure 1", "a-f", "Benchmark construction, data access, endpoint harmonization and module coverage", "Supplementary Tables 1, 2, 3, 6, 16"),
    ("Figure 2", "a-f", "Ecological modules, signed-rank score, optimizer, edges and aligned ablation space", "Supplementary Tables 7, 8, 9, 13"),
    ("Figure 3", "a-f", "Primary melanoma performance, signature family FDR, LODO and decision curve", "Supplementary Tables 10, 11, 12, 14"),
    ("Figure 4", "a-f", "Endpoint sensitivity, aligned locked-panel ablation, calibration and claim gate", "Supplementary Tables 6, 10, 11, 13, 14"),
    ("Figure 5", "a-f", "Locked external validation, panel transfer and PD1-like stress analysis", "Supplementary Tables 15, 16, 17"),
    ("Figure 6", "a-f", "Single-cell localization, ecological edges and perturbation hypotheses", "Supplementary Tables 7, 8, 18, 19"),
    ("Figure 7", "a-e", "Locked panel, thresholds, external scoring path and reproducibility boundary", "Supplementary Tables 7, 15, 16, 20"),
    ("Supplementary Figure 1", "a-b", "Cohort curation and access inventory", "Supplementary Tables 1, 2, 3"),
    ("Supplementary Figure 2", "a-b", "Platform and module gene coverage QC", "Supplementary Tables 3, 16"),
    ("Supplementary Figure 3", "a-b", "Discovery benchmark detailed performance", "Supplementary Tables 10, 12"),
    ("Supplementary Figure 4", "a", "Full baseline comparison", "Supplementary Table 10"),
    ("Supplementary Figure 5", "a", "Endpoint sensitivity", "Supplementary Table 6"),
    ("Supplementary Figure 6", "a-b", "Aligned locked-panel ablation and optimizer diagnostics", "Supplementary Tables 9, 13"),
    ("Supplementary Figure 7", "a", "Expanded external validation", "Supplementary Table 15"),
    ("Supplementary Figure 8", "a", "PD1-like stress rescue", "Supplementary Table 17"),
    ("Supplementary Figure 9", "a", "Single-cell and ecological mechanism", "Supplementary Table 18"),
    ("Supplementary Figure 10", "a-b", "Reproducibility and locked scoring package", "Supplementary Table 20"),
]

FIGURE_TO_TABLES = {
    "Fig1_source": [1, 2, 3, 6, 16],
    "Fig2_source": [7, 8, 9, 13],
    "Fig3_source": [10, 11, 12, 14],
    "Fig4_source": [6, 10, 11, 13, 14],
    "Fig5_source": [15, 16, 17],
    "Fig6_source": [7, 8, 18, 19],
    "Fig7_source": [7, 15, 16, 20],
    "SuppFig1_source": [1, 2, 3],
    "SuppFig2_source": [3, 16],
    "SuppFig3_source": [10, 12],
    "SuppFig4_source": [10],
    "SuppFig5_source": [6],
    "SuppFig6_source": [9, 13],
    "SuppFig7_source": [15],
    "SuppFig8_source": [17],
    "SuppFig9_source": [18],
    "SuppFig10_source": [20],
}


def table_path(n: int) -> Path:
    matches = sorted(TABLE_DIR.glob(f"supp_table_{n:02d}_*.tsv"))
    if not matches:
        raise FileNotFoundError(f"No supplementary table {n:02d} found in {TABLE_DIR}")
    return matches[0]


def safe_sheet_name(name: str, used: set[str]) -> str:
    cleaned = re.sub(r"[\[\]\*\?/\\:]", "_", name)[:31]
    base = cleaned
    i = 1
    while cleaned in used:
        suffix = f"_{i}"
        cleaned = f"{base[:31-len(suffix)]}{suffix}"
        i += 1
    used.add(cleaned)
    return cleaned


def read_table(n: int) -> pd.DataFrame:
    path = table_path(n)
    df = pd.read_csv(path, sep="\t")
    df.insert(0, "source_table", path.name)
    return df


def autosize_xlsx(path: Path) -> None:
    wb = load_workbook(path)
    header_fill = PatternFill("solid", fgColor="D9EAF7")
    for ws in wb.worksheets:
        ws.freeze_panes = "A2"
        for cell in ws[1]:
            cell.font = Font(bold=True)
            cell.fill = header_fill
        for col_idx, column_cells in enumerate(ws.columns, 1):
            max_len = 8
            for cell in column_cells[:200]:
                value = "" if cell.value is None else str(cell.value)
                max_len = max(max_len, min(len(value), 70))
            ws.column_dimensions[get_column_letter(col_idx)].width = min(max_len + 2, 48)
    wb.save(path)


def make_source_data() -> None:
    SUBMISSION_DIR.mkdir(parents=True, exist_ok=True)
    used: set[str] = set()
    with pd.ExcelWriter(SOURCE_DATA, engine="openpyxl") as writer:
        readme = pd.DataFrame(
            [
                {"field": "article", "value": "EcoNiche-Opt Communications Medicine submission"},
                {"field": "generated_on", "value": date.today().isoformat()},
                {"field": "repository", "value": REPO},
                {"field": "release_tag", "value": RELEASE_TAG},
                {"field": "release_archive", "value": ARCHIVE},
                {"field": "contents", "value": "Source data for main Figures 1-7, major Supplementary Figures 1-10, and Supplementary Tables 1-20."},
                {"field": "provenance", "value": "All sheets are generated from registered TSV outputs under tables/article and figure_manifest.tsv."},
            ]
        )
        readme.to_excel(writer, sheet_name=safe_sheet_name("README", used), index=False)

        pd.DataFrame(
            FIGURE_SOURCE_MAP,
            columns=["display_item", "panels", "content", "source_data_sheets"],
        ).to_excel(writer, sheet_name=safe_sheet_name("Figure_Source_Map", used), index=False)

        if (FIGURE_DIR / "figure_manifest.tsv").exists():
            pd.read_csv(FIGURE_DIR / "figure_manifest.tsv", sep="\t").to_excel(
                writer, sheet_name=safe_sheet_name("Figure_File_Manifest", used), index=False
            )
        if (TABLE_DIR / "table_manifest.tsv").exists():
            pd.read_csv(TABLE_DIR / "table_manifest.tsv", sep="\t").to_excel(
                writer, sheet_name=safe_sheet_name("Table_File_Manifest", used), index=False
            )

        for sheet, nums in FIGURE_TO_TABLES.items():
            frames = [read_table(n) for n in nums]
            combined = pd.concat(frames, ignore_index=True, sort=False)
            combined.to_excel(writer, sheet_name=safe_sheet_name(sheet, used), index=False)

        for n in range(1, 21):
            path = table_path(n)
            short = path.stem.replace("supp_table_", "ST")
            pd.read_csv(path, sep="\t").to_excel(writer, sheet_name=safe_sheet_name(short, used), index=False)
    autosize_xlsx(SOURCE_DATA)


def make_checklists() -> None:
    SUBMISSION_DIR.mkdir(parents=True, exist_ok=True)
    strobe = [
        (1, "Title and abstract", "Title and structured Abstract", "Article title and Abstract identify a multicohort biomarker/prediction study using retrospective public cohorts."),
        (2, "Background/rationale", "Introduction", "Rationale describes label, platform, cohort and immune-ecology heterogeneity."),
        (3, "Objectives", "Introduction final paragraph", "Objective is to develop and validate a locked immune-ecology transcriptomic score."),
        (4, "Study design", "Results Figure 1; Methods Data registration", "Retrospective multicohort benchmark with discovery, LODO, locked external and panel-transfer analyses."),
        (5, "Setting", "Methods Data registration; Supplementary Tables 1-3", "Public ICB transcriptomic cohorts with registered access status and cohort roles."),
        (6, "Participants", "Methods Response-label curation; Supplementary Tables 4-6", "Samples are included when expression, patient/sample mapping and endpoint labels are traceable."),
        (7, "Variables", "Methods Endpoint harmonization; Supplementary Table 6", "Response labels, cohort role, cancer type, therapy, time point, module coverage and model scores."),
        (8, "Data sources/measurement", "Methods Expression processing; Supplementary Tables 1-5", "Expression matrices and clinical annotations are taken from public datasets and source publications."),
        (9, "Bias", "Methods Locked validation; claim gate", "Leakage control, endpoint stratification, LODO testing and FDR-aware claim gating address selection and reporting bias."),
        (10, "Study size", "Results Figure 4a; Supplementary Table 6", "Sample counts are determined by endpoint eligibility and traceable expression/label availability."),
        (11, "Quantitative variables", "Methods Module score and statistical analysis", "Expression is rank-normalized and summarized as signed ecological module scores."),
        (12, "Statistical methods", "Methods Statistics and reproducibility", "AUROC, AUPRC, balanced accuracy, ECE, Brier score, decision curve, paired tests and FDR correction."),
        (13, "Participants/results flow", "Results Figures 1 and 4; Supplementary Tables 1-6", "Cohort registration and endpoint inclusion define the analysis flow."),
        (14, "Descriptive data", "Supplementary Tables 1-6", "Cohort roles, access status, expression QC and label evidence are tabulated."),
        (15, "Outcome data", "Supplementary Table 6", "Endpoint-specific responder/nonresponder counts are reported."),
        (16, "Main results", "Results Figures 3-5", "Primary melanoma, locked external and panel-transfer performance are reported with comparator families."),
        (17, "Other analyses", "Results Figures 4-6", "Endpoint sensitivity, ablation, calibration, single-cell and perturbation analyses."),
        (18, "Key results", "Discussion", "Evidence hierarchy emphasizes predeclared family-level improvement and refitting-free external scoring."),
        (19, "Limitations/boundary", "Discussion claim hierarchy", "Claim boundaries are defined through FDR-aware tiers and hypothesis-only mechanistic labels."),
        (20, "Interpretation", "Discussion", "Model is interpreted as coupled immune-ecology scoring rather than a single signature."),
        (21, "Generalisability", "Results Figure 5; Discussion", "External and panel-transfer cohorts evaluate refitting-free portability."),
        (22, "Funding", "Declarations/Funding", "The manuscript states that no specific funding was received for this work."),
    ]
    tripod = [
        (1, "Title/abstract", "Title and Abstract", "Prediction-model nature, setting and outcome are stated."),
        (2, "Background and objectives", "Introduction", "Need for a leakage-safe multicohort biomarker framework is stated."),
        (3, "Source of data", "Methods Data registration", "Retrospective public cohorts with registered access status."),
        (4, "Participants", "Methods Response-label curation", "Eligibility is based on traceable sample, patient, baseline/timepoint and endpoint mapping."),
        (5, "Outcome", "Methods Endpoint harmonization", "Response endpoints are strict RECIST, primary RECIST and clinical benefit."),
        (6, "Candidate predictors", "Methods Module score; Supplementary Table 7", "Signed-rank immune-ecology module scores and interaction edges."),
        (7, "Sample size", "Supplementary Table 6", "Endpoint-specific sample sizes are tabulated."),
        (8, "Missing data", "Methods Expression QC and gene coverage", "Module coverage and missing-gene handling are audited."),
        (9, "Statistical analysis methods", "Methods Statistics and reproducibility", "Optimization, calibration, thresholds, comparator tests and FDR correction are specified."),
        (10, "Risk groups/thresholds", "Results Figure 7; Methods Thresholding", "Endpoint-specific discovery thresholds are locked before external scoring."),
        (11, "Model development", "Methods Model objective", "Signed-rank module score, interaction terms and biologically constrained objective are defined."),
        (12, "Internal validation", "Results Figure 3; Supplementary Table 12", "Leave-one-dataset-out validation reports holdout AUROC."),
        (13, "External validation", "Results Figure 5; Supplementary Table 15", "Locked external/panel cohorts are scored without refitting."),
        (14, "Model specification", "Methods; Code availability", "Formula, package API, CLI commands and source archive are provided."),
        (15, "Model performance", "Results Figures 3-5", "AUROC, AUPRC, balanced accuracy, ECE, Brier score and decision-curve metrics."),
        (16, "Model updating", "Results Figure 5e-f", "PD1-like transfer-head analysis is separately labelled from the locked primary model."),
        (17, "Interpretation", "Discussion", "The score is interpreted through immune-response and resistance ecological states."),
        (18, "Implications", "Discussion; Figure 7", "Locked 62-gene panel-compatible scoring rule enables independent computation."),
        (19, "Supplementary resources", "Supplementary Information; Source Data", "Supplementary tables, source data workbook and repository files support reproducibility."),
        (20, "Protocol/registration", "Supplementary Tables 1-5", "Cohort registration and source evidence are documented."),
    ]
    remark = [
        (1, "Marker/objective", "Title, Abstract, Introduction", "EcoNiche-Opt is a transcriptomic immune-ecology biomarker score for ICB response."),
        (2, "Patient/source material", "Methods Data registration; Supplementary Tables 1-5", "Public tumour expression cohorts and clinical response annotations."),
        (3, "Clinical endpoints", "Methods Endpoint harmonization; Supplementary Table 6", "Strict RECIST, primary RECIST and clinical benefit endpoints."),
        (4, "Assay method", "Methods Expression processing", "RNA-seq, microarray or targeted-expression matrices are harmonized into rank-normalized module scores."),
        (5, "Marker definition", "Methods Module score; Figure 2", "Signed module scores, ecological interaction edges and locked 62-gene panel."),
        (6, "Study design", "Figure 1; Methods", "Multicohort discovery, LODO, locked external and panel-transfer design."),
        (7, "Statistical analysis", "Methods Statistics and reproducibility", "Performance metrics, paired tests, FDR correction, calibration and decision curves."),
        (8, "Cutpoint definition", "Methods Thresholding; Figure 7", "Endpoint thresholds selected only in discovery cohorts and then frozen."),
        (9, "Missing data/coverage", "Methods Gene coverage; Supplementary Table 16", "Coverage fraction evaluates computability across platforms."),
        (10, "Results and estimates", "Results Figures 3-5", "Effect estimates include AUROC, deltas, confidence intervals and q values where available."),
        (11, "Validation", "Results Figure 5", "Locked external and panel-transfer cohorts evaluate portability without refitting."),
        (12, "Biological interpretation", "Results Figure 6", "Single-cell localization and ecological edges connect the score to immune niches."),
        (13, "Data/code sharing", "Data availability; Code availability", "Public repository, release archive and source data workbook are supplied."),
        (14, "Discussion", "Discussion", "Claim hierarchy separates primary benchmark evidence, external support and mechanistic hypotheses."),
    ]
    reporting = [
        ("Statistics", "Sample sizes", "Endpoint-specific sample sizes are reported in Figure 4a and Supplementary Table 6."),
        ("Statistics", "Replicates", "Biological units are patient/sample-level public cohort specimens; LODO uses dataset-level holdouts."),
        ("Statistics", "Randomization", "No treatment randomization was performed by this retrospective computational study."),
        ("Statistics", "Blinding", "Locked external scoring computes predictions before external labels are used for metric calculation."),
        ("Statistics", "Statistical tests", "Paired bootstrap/DeLong-compatible comparisons and Benjamini-Hochberg FDR correction are described."),
        ("Data", "Data availability", "Public and restricted datasets are listed with access status; source data workbook is provided."),
        ("Code", "Code availability", f"Repository {REPO}; release tag {RELEASE_TAG}; release archive {ARCHIVE}."),
        ("Software", "Core software", "Python package econiche_opt v0.3.1; R wrapper via reticulate; scripts listed in Code availability."),
        ("Reporting", "Guidelines", "STROBE, TRIPOD and REMARK checklists are supplied for the observational biomarker prediction study."),
    ]
    with pd.ExcelWriter(CHECKLISTS, engine="openpyxl") as writer:
        pd.DataFrame(reporting, columns=["section", "topic", "response"]).to_excel(
            writer, sheet_name="Nature_Reporting_Summary", index=False
        )
        for sheet, rows in {"STROBE": strobe, "TRIPOD": tripod, "REMARK": remark}.items():
            pd.DataFrame(rows, columns=["item", "topic", "manuscript_location", "response"]).to_excel(
                writer, sheet_name=sheet, index=False
            )
    autosize_xlsx(CHECKLISTS)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Arial"


def make_reporting_summary_docx() -> None:
    SUBMISSION_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Nature Portfolio Reporting Summary\nEcoNiche-Opt")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(16)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"Prepared for Communications Medicine submission | Generated {date.today().isoformat()}")

    sections = {
        "Study design": [
            "Retrospective multicohort transcriptomic biomarker and prediction-model study using public immune checkpoint blockade cohorts.",
            "Discovery, leave-one-dataset-out validation, locked external scoring and panel-transfer analyses are kept separate by design.",
        ],
        "Sample definition and endpoints": [
            "Eligible samples require traceable expression data, sample/patient mapping and response-label evidence.",
            "Endpoints are strict RECIST, primary RECIST and clinical benefit; endpoint-specific sample sizes are reported in Supplementary Table 6.",
        ],
        "Statistics and reproducibility": [
            "Performance metrics include AUROC, AUPRC, balanced accuracy, Brier score, expected calibration error and decision-curve net benefit.",
            "Comparator analyses use paired bootstrap or DeLong-compatible tests with Benjamini-Hochberg false-discovery-rate correction.",
            "All training, feature selection, thresholding and calibration are performed only within discovery/training data.",
        ],
        "Data availability": [
            "Public data sources and access status are reported in Supplementary Tables 1-6.",
            "Source Data.xlsx contains figure-level source mappings, figure/table manifests, and Supplementary Tables 1-20.",
            "Controlled or access-restricted cohorts are described by access route and are not redistributed.",
        ],
        "Code availability": [
            f"Repository: {REPO}",
            f"Manuscript release tag: {RELEASE_TAG}",
            f"Release-specific source archive: {ARCHIVE}",
            "Core commands include python -m pip install -e ., python -m pytest -q, python -m econiche_opt.cli make-demo, and python -m econiche_opt.cli validate-project --mode demo.",
        ],
        "Reporting guidelines": [
            "STROBE is supplied for the retrospective observational cohort structure.",
            "TRIPOD is supplied for prediction-model development and validation reporting.",
            "REMARK is supplied for oncology biomarker reporting.",
        ],
    }
    for heading, bullets in sections.items():
        add_heading(doc, heading, 1)
        for text in bullets:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(text)
    doc.save(REPORTING_SUMMARY)


def make_readiness_note() -> None:
    title = "EcoNiche-Opt: a locked immune-ecology transcriptomic score for multicohort prediction of immune checkpoint blockade response"
    word_count = len(re.findall(r"[A-Za-z0-9]+(?:[-'][A-Za-z0-9]+)*", title))
    lines = [
        "# Communications Medicine submission readiness",
        "",
        f"- Generated: {date.today().isoformat()}",
        f"- Title word count: {word_count} words (Communications Medicine guide: 15 words or fewer).",
        "- Abstract: structured as Background, Methods, Results and Conclusions; 250-word limit checked in the manuscript source.",
        "- Plain language summary: added immediately after the Abstract and kept near the 120-word guide.",
        "- Transparent peer review: cover letter states opt-in.",
        f"- Code availability: repository, release tag and release-specific archive are listed ({ARCHIVE}).",
        "- Source data: Source_Data.xlsx covers Figures 1-7, Supplementary Figures 1-10 and Supplementary Tables 1-20.",
        "- Reporting files: Nature Portfolio Reporting Summary document plus STROBE, TRIPOD and REMARK checklist workbook.",
        "",
        "Files:",
        f"- {SOURCE_DATA.relative_to(ROOT)}",
        f"- {REPORTING_SUMMARY.relative_to(ROOT)}",
        f"- {CHECKLISTS.relative_to(ROOT)}",
    ]
    READINESS.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    make_source_data()
    make_checklists()
    make_reporting_summary_docx()
    make_readiness_note()
    print(f"Wrote {SOURCE_DATA}")
    print(f"Wrote {REPORTING_SUMMARY}")
    print(f"Wrote {CHECKLISTS}")
    print(f"Wrote {READINESS}")


if __name__ == "__main__":
    main()

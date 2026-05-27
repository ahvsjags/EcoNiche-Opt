from __future__ import annotations

import re
import shutil
import sys
import zipfile
from collections import OrderedDict
from pathlib import Path

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from scripts.documents.make_word_from_manuscript import (  # noqa: E402
    EQUATIONS,
    FIG_DIR,
    MAIN_FIGURES_HIRES,
    add_runs_from_markdown,
)


PAPER = ROOT / "paper"
SOURCE_MAIN_MD = PAPER / "econiche_opt_manuscript_en_v1_20260509.md"
SOURCE_SUPP_DOCX = PAPER / "communications_medicine_submission" / "EcoNiche-Opt_Supplementary_Information.docx"
SOURCE_SUPP_PDF = PAPER / "communications_medicine_submission" / "EcoNiche-Opt_Supplementary_Information.pdf"
SOURCE_DATA = PAPER / "communications_medicine_submission" / "Source_Data.xlsx"
SOURCE_CHECKLIST = PAPER / "communications_medicine_submission" / "STROBE_TRIPOD_REMARK_checklists.xlsx"

OUT_DIR = PAPER / "Journal of Translational Medicine投稿"
BUILD_DIR = OUT_DIR / "_build"
EQ_DIR = BUILD_DIR / "equations_600dpi"

TITLE = "EcoNiche-Opt: a locked immune-ecology transcriptomic score for multicohort prediction of immune checkpoint blockade response"
AUTHORS = "Pengyuan Xu1,3, Guang Yang2,3, Moyan Li3*"
AFFILIATIONS = [
    "1 Department of Materials Science and Engineering, Monash University, Clayton, VIC 3800, Australia.",
    "2 School of Economics and Management, China University of Mining and Technology, Xuzhou 221116, Jiangsu, China.",
    "3 Hong Kong University of Science and Technology (Guangzhou), Guangzhou 510000, China.",
    "* Correspondence: moyanli@hkust-gz.edu.cn",
]
KEYWORDS = (
    "immune checkpoint blockade; transcriptomic biomarker; melanoma; tumour immune microenvironment; "
    "machine learning; external validation; immune ecology; precision oncology"
)

MAIN_MD_OUT = OUT_DIR / "EcoNiche-Opt_JTM_Main_Manuscript.md"
MAIN_DOCX_OUT = OUT_DIR / "EcoNiche-Opt_JTM_Main_Manuscript.docx"
COVER_MD_OUT = OUT_DIR / "EcoNiche-Opt_JTM_Cover_Letter.md"
COVER_DOCX_OUT = OUT_DIR / "EcoNiche-Opt_JTM_Cover_Letter.docx"
README_OUT = OUT_DIR / "JTM_submission_upload_notes.md"


def split_sections(lines: list[str]) -> OrderedDict[str, list[str]]:
    sections: OrderedDict[str, list[str]] = OrderedDict()
    current: str | None = None
    for line in lines:
        if line.startswith("## "):
            current = line[3:].strip()
            sections[current] = []
            continue
        if current is not None:
            sections[current].append(line.rstrip())
    return sections


def clean_block(block: list[str]) -> list[str]:
    cleaned: list[str] = []
    for line in block:
        if line.strip() in {"[Author information pending]", ""}:
            cleaned.append(line.rstrip())
            continue
        cleaned.append(line.rstrip())
    while cleaned and not cleaned[0].strip():
        cleaned.pop(0)
    while cleaned and not cleaned[-1].strip():
        cleaned.pop()
    return cleaned


def bmc_main_markdown() -> str:
    lines = SOURCE_MAIN_MD.read_text(encoding="utf-8").splitlines()
    sections = split_sections(lines)
    abstract = clean_block(sections["Abstract"])
    background = clean_block(sections["Introduction"])
    methods = clean_block(sections["Methods"])
    results = clean_block(sections["Results"])
    discussion = clean_block(sections["Discussion"])
    data_availability = clean_block(sections["Data availability"])
    code_availability = clean_block(sections["Code availability"])
    references = clean_block(sections["References"])
    figure_legends = clean_block(sections["Figure Legends"])

    output: list[str] = [
        f"# {TITLE}",
        "",
        AUTHORS,
        "",
        *AFFILIATIONS,
        "",
        "## Abstract",
        "",
        *abstract,
        "",
        "## Keywords",
        "",
        KEYWORDS,
        "",
        "## Background",
        "",
        *background,
        "",
        "## Methods",
        "",
        *methods,
        "",
        "## Results",
        "",
        *results,
        "",
        "## Discussion",
        "",
        *discussion,
        "",
        "## Conclusions",
        "",
        (
            "EcoNiche-Opt reframes immune-checkpoint blockade transcriptomic prediction as an auditable "
            "immune-ecology biomarker framework rather than another isolated response signature. By combining "
            "traceable public-cohort curation, endpoint harmonization, signed-rank ecological module scoring, "
            "interaction-edge modelling, claim-gated multicohort benchmarking, locked external scoring, and a "
            "62-gene panel-compatible implementation, the framework provides a reproducible route from heterogeneous "
            "ICB cohorts to independently computable biomarker scores."
        ),
        "",
        "## Abbreviations",
        "",
        (
            "APM, antigen-presentation machinery; AUROC, area under the receiver operating characteristic curve; "
            "AUPRC, area under the precision-recall curve; DCB, durable clinical benefit; ECE, expected calibration "
            "error; ICB, immune checkpoint blockade; LODO, leave-one-dataset-out; RECIST, Response Evaluation Criteria "
            "in Solid Tumours; SAP, statistical analysis plan."
        ),
        "",
        "## Declarations",
        "",
        "### Ethics approval and consent to participate",
        "",
        (
            "This study used publicly available or controlled-access de-identified transcriptomic and clinical "
            "annotation data from previously published studies. No new human participants were recruited and no new "
            "biospecimens were collected. Ethics approvals and consent procedures for the original cohorts were "
            "reported by the source studies, and controlled-access datasets remain subject to the original repository "
            "and data-use conditions."
        ),
        "",
        "### Consent for publication",
        "",
        "Not applicable.",
        "",
        "### Availability of data and materials",
        "",
        *data_availability,
        "",
        "### Availability of code",
        "",
        *code_availability,
        "",
        "### Competing interests",
        "",
        "The authors declare that they have no competing interests.",
        "",
        "### Funding",
        "",
        "The authors declare that no specific funding was received for this work.",
        "",
        "### Authors' contributions",
        "",
        (
            "PX and ML conceived the study. PX implemented the computational analyses, model packaging, data curation "
            "workflow, benchmark evaluation, figures, and tables. GY contributed to statistical interpretation, "
            "translational framing, and manuscript revision. ML supervised the study, interpreted results, revised the "
            "manuscript, and serves as corresponding author. All authors read and approved the final manuscript."
        ),
        "",
        "### Acknowledgements",
        "",
        "Not applicable.",
        "",
        "## References",
        "",
        *references,
        "",
        "## Figure legends",
        "",
        *figure_legends,
        "",
    ]
    return "\n".join(output).replace("## Figure Legends", "## Figure legends")


def configure_bmc_doc(doc: Document, line_numbers: bool = True, page_numbers: bool = True) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    ln = section._sectPr.find(qn("w:lnNumType"))
    if ln is not None:
        section._sectPr.remove(ln)
    if line_numbers:
        ln = OxmlElement("w:lnNumType")
        ln.set(qn("w:countBy"), "1")
        ln.set(qn("w:restart"), "continuous")
        section._sectPr.append(ln)

    if page_numbers:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("Page ")
        field_begin = OxmlElement("w:fldChar")
        field_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        field_end = OxmlElement("w:fldChar")
        field_end.set(qn("w:fldCharType"), "end")
        run._r.append(field_begin)
        run._r.append(instr)
        run._r.append(field_end)

    styles = doc.styles
    for name, size, bold in [
        ("Normal", 10, False),
        ("Heading 1", 14, True),
        ("Heading 2", 12, True),
        ("Heading 3", 11, True),
        ("Title", 16, True),
    ]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor(0, 0, 0)


def render_equation(eq_num: int) -> Path:
    EQ_DIR.mkdir(parents=True, exist_ok=True)
    out = EQ_DIR / f"equation_{eq_num:02d}.png"
    lines = EQUATIONS[eq_num]
    fig_height = 0.48 + 0.28 * (len(lines) - 1)
    fig = plt.figure(figsize=(5.8, fig_height), dpi=600, facecolor="white")
    ax = fig.add_axes([0, 0, 1, 1])
    ax.axis("off")
    if len(lines) == 1:
        y_positions = [0.5]
    elif len(lines) == 2:
        y_positions = [0.64, 0.34]
    else:
        y_positions = [0.73, 0.5, 0.27][: len(lines)]
    font_size = 10 if eq_num not in {4, 6, 7, 8, 10, 12} else 9.3
    for line, y in zip(lines, y_positions):
        ax.text(0.47, y, f"${line}$", ha="center", va="center", fontsize=font_size)
    ax.text(0.985, 0.5, f"({eq_num})", ha="right", va="center", fontsize=9.5)
    fig.savefig(out, dpi=600, transparent=False, bbox_inches="tight", pad_inches=0.03)
    plt.close(fig)
    return out


def add_page_number_update_setting(docx_path: Path) -> None:
    settings_name = "word/settings.xml"
    tmp_path = docx_path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == settings_name:
                xml = data.decode("utf-8")
                if "doNotCompressPictures" not in xml:
                    xml = xml.replace(
                        "</w:settings>",
                        '<w:doNotCompressPictures/><w:defaultImageDpi w:val="600"/></w:settings>',
                    )
                data = xml.encode("utf-8")
            zout.writestr(item, data)
    tmp_path.replace(docx_path)


def add_markdown_line(doc: Document, raw: str) -> None:
    line = raw.rstrip()
    if not line.strip():
        return
    if line.startswith("# "):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_runs_from_markdown(p, line[2:].strip(), bold_default=True)
        for run in p.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.space_after = Pt(12)
        return
    if line.startswith("## "):
        p = doc.add_paragraph(style="Heading 1")
        add_runs_from_markdown(p, line[3:].strip(), bold_default=True)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        return
    if line.startswith("### "):
        p = doc.add_paragraph(style="Heading 2")
        add_runs_from_markdown(p, line[4:].strip(), bold_default=True)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        return
    match = re.match(r"^Formula\s+(\d+):\s*`.*`\s*$", line)
    if match:
        eq_num = int(match.group(1))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(6)
        p.add_run().add_picture(str(render_equation(eq_num)), width=Inches(5.4))
        return
    if line.startswith("- "):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 2.0
        p.paragraph_format.space_after = Pt(3)
        add_runs_from_markdown(p, line[2:].strip())
        return
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_after = Pt(0)
    add_runs_from_markdown(p, line)


def markdown_to_docx(markdown: str, out_path: Path, line_numbers: bool = True, page_numbers: bool = True) -> None:
    doc = Document()
    configure_bmc_doc(doc, line_numbers=line_numbers, page_numbers=page_numbers)
    for line in markdown.splitlines():
        add_markdown_line(doc, line)
    doc.core_properties.title = TITLE
    doc.core_properties.subject = "Journal of Translational Medicine submission"
    doc.core_properties.author = "EcoNiche-Opt authors"
    doc.save(out_path)
    add_page_number_update_setting(out_path)


def make_cover_letter() -> str:
    return f"""# Cover letter

Dear Editors,

We submit "{TITLE}" for consideration as a Research Article in Journal of Translational Medicine, with strongest relevance to the Disease Biomarkers, Cancer Microenvironment, and Data-driven Clinical Decision Processes sections.

The study addresses a translational barrier in immune-checkpoint blockade biomarker development: public transcriptomic cohorts contain useful immune signals, but response labels, sampling status, assay platforms, endpoint definitions, and tumour-immune ecological states are often heterogeneous. EcoNiche-Opt addresses this gap by combining auditable cohort and response-label curation, endpoint-stratified evaluation, signed-rank ecological module scoring, module-interaction edges, biologically constrained optimization, leave-one-dataset-out testing, locked external validation, and a reproducible 62-gene panel-compatible scoring rule.

In primary melanoma benchmarks, EcoNiche-Opt showed FDR-supported improvement over a predeclared eight-signature family. Locked external and panel-transfer analyses showed reproducible portability without model refitting, and single-cell localization plus ecological-edge analyses linked the score to antigen presentation, T/NK effector activity, T-cell dysfunction, myeloid suppression, and stromal exclusion. The work therefore provides not only a predictor, but a reusable translational framework for developing and validating immune-ecology biomarkers across heterogeneous ICB cohorts.

All authors have approved the manuscript for submission. The manuscript has not been published and is not under consideration elsewhere. The authors declare no competing interests. The study used previously published de-identified public or controlled-access data and did not recruit new human participants or collect new biospecimens.

Sincerely,

Moyan Li, corresponding author

Hong Kong University of Science and Technology (Guangzhou), Guangzhou 510000, China

moyanli@hkust-gz.edu.cn
"""


def copy_submission_assets() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    BUILD_DIR.mkdir(parents=True, exist_ok=True)
    for src, dst_name in [
        (SOURCE_SUPP_DOCX, "Additional_file_1_Supplementary_Information.docx"),
        (SOURCE_SUPP_PDF, "Additional_file_1_Supplementary_Information.pdf"),
        (SOURCE_DATA, "Additional_file_2_Source_Data.xlsx"),
        (SOURCE_CHECKLIST, "Additional_file_3_STROBE_TRIPOD_REMARK_checklists.xlsx"),
    ]:
        shutil.copy2(src, OUT_DIR / dst_name)
    for idx, fig_name in MAIN_FIGURES_HIRES.items():
        shutil.copy2(FIG_DIR / fig_name, OUT_DIR / f"Figure_{idx}.png")


def make_readme() -> str:
    return """# Journal of Translational Medicine submission upload notes

Recommended article type: Research Article.

Recommended section: Disease Biomarkers. If the system allows a secondary section, use Cancer Microenvironment or Data-driven Clinical Decision Processes.

Upload files:

- Main manuscript: EcoNiche-Opt_JTM_Main_Manuscript.docx
- Cover letter: EcoNiche-Opt_JTM_Cover_Letter.docx
- Additional file 1: Additional_file_1_Supplementary_Information.pdf
- Source data / additional file: Additional_file_2_Source_Data.xlsx
- Reporting checklists: Additional_file_3_STROBE_TRIPOD_REMARK_checklists.xlsx
- Figures: Figure_1.png through Figure_7.png

Do not upload NaturePortfolio_Reporting_Summary_EcoNicheOpt.docx for Journal of Translational Medicine unless the submission system specifically requests it; this file was prepared for Nature Portfolio submission.

Suggested keywords: immune checkpoint blockade; transcriptomic biomarker; melanoma; tumour immune microenvironment; machine learning; external validation; immune ecology; precision oncology.
"""


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    copy_submission_assets()
    main_md = bmc_main_markdown()
    MAIN_MD_OUT.write_text(main_md, encoding="utf-8")
    markdown_to_docx(main_md, MAIN_DOCX_OUT)
    cover_md = make_cover_letter()
    COVER_MD_OUT.write_text(cover_md, encoding="utf-8")
    markdown_to_docx(cover_md, COVER_DOCX_OUT, line_numbers=False, page_numbers=False)
    README_OUT.write_text(make_readme(), encoding="utf-8")
    print(MAIN_DOCX_OUT)
    print(COVER_DOCX_OUT)
    print(README_OUT)


if __name__ == "__main__":
    main()
